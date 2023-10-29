import os
import re
import time
import json
import docx
import joblib

import shutil
import psycopg2
from pathlib import Path
from loguru import logger  
from dotenv import load_dotenv

import pandas as pd
import numpy as np 

import nltk
from nltk.corpus import stopwords
from pymystem3 import Mystem
from string import punctuation
from gensim.models import KeyedVectors

from fastapi.middleware.cors import CORSMiddleware
from fastapi import FastAPI, File, Request, UploadFile, Response
from fastapi.responses import JSONResponse

from sklearn.metrics.pairwise import linear_kernel
from sklearn.feature_extraction.text import CountVectorizer, TfidfVectorizer
from sklearn.metrics.pairwise import linear_kernel,cosine_similarity

nltk.download('stopwords')

# load labels list
labels = 0
labels_path = Path().cwd().joinpath('names').joinpath('lables.json')
with open(labels_path, 'rb') as f:
    labels = json.load(f)
labels = labels['data']

# Load the model clf
clf_path = Path().cwd().joinpath('names').joinpath('models').joinpath('svm_model.pkl')
clf = joblib.load(clf_path.__str__())

# Load the KeyedVectors model
word2vec_model_path = Path().cwd().joinpath('names').joinpath('models').joinpath('word2vec_model.model')
word2vec_model = KeyedVectors.load(word2vec_model_path.__str__())

# Load info about GOSTS
standars_info_path = Path().cwd().joinpath('gosts').joinpath('standarts.csv')
standars_info = pd.read_csv(standars_info_path)

# Файл для сопоставления ГОСТОВ и Групп\Подгрупп продукции
gost_compare = 0
gost_compare_path = Path().cwd().joinpath("gosts").joinpath("gosts_compare.json")
with open(gost_compare_path, 'rb') as f:
    gost_compare = json.load(f)
gost_compare = gost_compare["data"]

our_data=pd.read_csv(Path.cwd().joinpath("data_classes.csv"), encoding="utf-8", delimiter='"')
mystem = Mystem() 
russian_stopwords = stopwords.words("russian")
nltk.download("stopwords")

# Стандартная функция обработки естественного языка
def clean_text(text):
    text = re.sub(r"[^а-яА-Я\s]", "", text)  # Remove all characters except letters and spaces
    text = text.lower()  # Convert text to lowercase
    text = text.split()  # Split the text into words
    return text

# Предсказание класса для подзадачи НАИМЕНОВАНИЕ - НАГДРУППА ПРОДУКЦИИ
def predict_class(input_text):
    cleaned_text = clean_text(input_text)
    vector = np.mean([word2vec_model[word] for word in cleaned_text if word in word2vec_model], axis=0)
    predicted_class = clf.predict([vector])
    return predicted_class[0]

# Предстаказание класса по второму методу
def predict_text(text, model_file):
    # Загрузить сохраненный пайплайн из файла
    pipeline = joblib.load(model_file)
    
    # Вызвать метод predict на пайплайне для предсказания класса текста
    predicted_class = pipeline.predict([text])[0]
    
    return predicted_class


# Функция для метрики обоснованности выбора по каждому из оборудований
def compute_similarity_precentages(sentences1, sentences2):
    all_sentences = sentences1 + sentences2
    tfidf_vectorizer = CountVectorizer(ngram_range=(1, 2))
    tfidf_matrix = tfidf_vectorizer.fit_transform(all_sentences)
    tfidf_matrix_sentences1 = tfidf_matrix[:len(sentences1)]
    tfidf_matrix_sentences2 = tfidf_matrix[len(sentences1):]
    similarity_matrix = cosine_similarity(tfidf_matrix_sentences1, tfidf_matrix_sentences2)
    
    similarity_percentages = [np.round(similarity * 100, 2) for similarity in similarity_matrix]

    return similarity_percentages 

# Метрика общей схожести всего оборудования с ГОСТами
def calculate_similarity_score(list1, list2):
    tfidf_vectorizer = CountVectorizer(ngram_range=(1, 2))

    tfidf_matrix_input = tfidf_vectorizer.fit_transform(list1)
    tfidf_matrix_target = tfidf_vectorizer.transform(list2)

    cosine_similarities = linear_kernel(tfidf_matrix_target, tfidf_matrix_input)

    return(np.mean(cosine_similarities))

# Вспомогательная функция для текстовой обработки
def preprocess_text(text):
    tokens = mystem.lemmatize(text.lower())
    tokens = [token for token in tokens if token not in russian_stopwords\
              and token != " " \
              and token.strip() not in punctuation]
    
    text = " ".join(tokens)
    
    return text

# Функция для реализации работы соотвествия ГОСТОв и оборудования в них
def get_predictions(input_strings,target_strings,similarity_threshold=0.5):
    tfidf_vectorizer = TfidfVectorizer(ngram_range=(1, 2))

    # Преобразуем тексты в матрицу TF-IDF
    tfidf_matrix_input = tfidf_vectorizer.fit_transform(input_strings)
    tfidf_matrix_target = tfidf_vectorizer.transform(target_strings)

    # Вычисляем косинусное сходство между входными и целевыми строками
    cosine_similarities = linear_kernel(tfidf_matrix_target, tfidf_matrix_input)

    # Находим индексы строк с сходством выше порога
    most_similar_indices = np.argwhere(cosine_similarities > similarity_threshold)

    # Выводим только строки, у которых сходство больше 70%
    list_values={}
    for i, j in most_similar_indices:
        input_str = input_strings[j]
        target_str = target_strings[i]
        similarity = cosine_similarities[i, j]
        list_values[input_str]=similarity
        # logger.debug(f"Целевая строка: {target_str}")
        # logger.debug(f"Наиболее подходящая строка: {input_str}")
        # logger.debug(f"Сходство (TF-IDF): {similarity}")
        # logger.debug("---------")
    return set(list_values)

# Далее идет инициализация Backend-части + работы с БД

PRODUCTION = False

DBenv = Path().cwd().parent.joinpath("DB.env")
load_dotenv(DBenv, override=True)

conn = 0
cur = 0

logger.debug("Waiting for DB service Up...")
time.sleep(5)

try:     
    HOST = None
    DBNAME = None
    PASSWORD = None

    if PRODUCTION:
        HOST=os.environ.get("DB_HOST")  
        DBNAME=os.environ.get("POSTGRES_DB")
        USER=os.environ.get("POSTGRES_USER")
        PASSWORD=os.environ.get("POSTGRES_PASSWORD")
    else:
        HOST=os.environ.get("DB_LOCAL")
        DBNAME=os.environ.get("LOCAL_DB")
        USER=os.environ.get("LOCAL_USER")
        PASSWORD=os.environ.get("LOCAL_PASSWORD")
    
    PORT=os.environ.get("PORT")

    logger.success(('Docker DB connection started \n', HOST, PORT, DBNAME, USER, PASSWORD, ' - env variables!'))

    conn = psycopg2.connect(
        dbname=DBNAME, 
        host=HOST, 
        user=USER, 
        password=PASSWORD, 
        port=PORT)

    cur = conn.cursor()

    if PRODUCTION:
        logger.success('Docker DB connected!')
    else:
        logger.success('Local DB connected!')

except Exception as e:
    logger.error(f'Database connect failed \n {e}!')

app = FastAPI()

origins = ["*"]

app.add_middleware(
    CORSMiddleware,
    allow_origins=origins,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

@app.get("/")
def root():
    return {"message": "Hello World"}


@app.get("/allGosts")
async def getAllGosts():
    """
    Вывод названий отработанных ГОСТов
    """
    cur.execute("SELECT * FROM gosts")
    data = cur.fetchall()
    
    gosts = []
    for index, subdata in enumerate(data):
        # logger.debug(subdata[1]) # name
        # logger.debug(subdata[2]) # equip
        gosts.append(subdata[1].replace('.docx', ''))

    return JSONResponse(content={"docs": gosts}, status_code=200)

@app.post("/get_pred")
async def upload_folder_get_pred(request: Request,  file: UploadFile = File(...)):
    """
    Функция для выявления оборудования согласно поданному файлу ГОСТа (внимание, здесь используется multipart-formdata!)
    """
    # Get the folder name from the request parameter
    path_to_save=Path.cwd().joinpath('download').joinpath(file.filename)
    with open(path_to_save, "wb") as dest_file:
        shutil.copyfileobj(file.file, dest_file)

    # Work on it
    doc = docx.Document(Path(path_to_save))
    full_text=[]
    for elem in doc.paragraphs:
        full_text.append(elem.text)

    prediction=get_predictions(our_data.text.values.tolist(),full_text)

    return {"prediction": prediction}

@app.post("/check_dataset")
async def upload_dataset(request: Request, file: UploadFile = File(...)):
    """
    Функция для загрузки датасета и проверки его в нашем ПК
    """
    path_to_save = Path.cwd().joinpath('datacheck').joinpath(file.filename)

    with open(path_to_save, "wb") as dest_file:
        shutil.copyfileobj(file.file, dest_file)
    
    cur.execute("SELECT * FROM gosts;")
    data = cur.fetchall()

    total_gost_data = []
    for index, subdata in enumerate(data):
        # logger.debug(subdata[0]) # id
        # logger.debug(subdata[1]) # name
        # logger.debug(subdata[2]) # equip

        total_gost_data.append({
            "name" : subdata[1],
            "equip": subdata[2]
        })

    save_df = pd.DataFrame(columns=["id", "Найденные вложенные ГОСТ", "Группа продукции", "Наименование продукции", "Коды ОКПД 2 / ТН ВЭД ЕАЭС", "Заявленное ТО", "Эталонное ТО (из ГОСТ)", "Оценка схожести"])
    result = []

    # Open via pandas
    df = pd.read_csv(path_to_save)
    for row in df.iterrows():
        index = row[1][0] # index
        gost = row[1][1] # GOSTS
        group = row[1][2] # Group
        name = row[1][3] # Naming
        tnved = row[1][4] # TN VED
        equip = row[1][5] # Apparats

        uppergroup_on_current_group = []

        # try to find in uppergroups
        for uppergroup in gost_compare:
            key = list(uppergroup.keys())[0]
            values = list(uppergroup.values())[0]

            if group in values:
                uppergroup_on_current_group.append(key)

        # try to find gost on uppergroups 
        uppergroup_gosts = [] 
        for subgroup in uppergroup_on_current_group:
            subgroup_gosts = standars_info[standars_info["Группа продукции"] == subgroup]["Обозначение и наименование стандарта"].to_list()
            uppergroup_gosts.extend(subgroup_gosts)
        uppergroup_gosts = list(set(uppergroup_gosts))

        # check with our table && compare search result to validate equipment
        find_gosts = []
        find_equipment = []

        # if find groups
        if uppergroup_gosts:
            for subdata in total_gost_data:
                if subdata["name"].replace(' ', '').replace('"', '').replace("'", '').replace('\xa0', '') in [x.replace(' ', '').replace('"', '').replace("'", '').replace('\xa0', '') for x in uppergroup_gosts]:
                    find_gosts.append(subdata["name"].replace('\xa0', ' '))
                    find_equipment.extend(subdata["equip"].split(';;'))
        
        similarity_score = 0

        try:
            # Обший процент
            similarity_score = calculate_similarity_score(find_equipment, equip.split(';'))
        except Exception as e:
            logger.error(e)

        # save_df = pd.concat(
        #     [
        #         save_df, 
        #         pd.DataFrame.from_dict(
        #             list(
        #                     {
        #                         "id": index,
        #                         "docs": find_gosts,
        #                         "group": group,
        #                         "name": name,
        #                         "tnved": tnved,
        #                         "equipment_user": equip.split(';'),
        #                         "equipment_find": find_equipment,
        #                         "similarity_score": similarity_score
        #                     }.items()
        #                 ), columns=["id", "docs", "group", "name", "tnved", "equipment_user", "equipment_find", "similarity_score"])], ignore_index=True)

        save_df.loc[len(save_df)] = [index, find_gosts, group, name, tnved, equip.split(';'), find_equipment, similarity_score]

        result.append(
            {
                "id": index,
                "docs": find_gosts,
                "group": group,
                "name": name,
                "tnved": tnved,
                "equipment_find_len": len(find_equipment),
                "equipment_find": find_equipment,
                "equipment_user": equip.split(';'),
                "equipment_user_len": len(equip.split(';')),
                "similarity_score": similarity_score
            }
        )

    save_df.to_csv('result.csv')
    
    # Прислать ему колонки и проверенный датасет
    return JSONResponse(status_code=200, content={"data": result})

@app.post("/train")
async def train_upload_dataset(request: Request, file: UploadFile = File(...)):
    """
    Функция для отработки всех данных
    """
    database = Path.cwd().joinpath('download')
    path_to_save = database.joinpath(file.filename)

    # saved file
    with open(path_to_save, "wb") as dest_file:
        shutil.copyfileobj(file.file, dest_file)

    # clear all database
    cur.execute('TRUNCATE gosts CASCADE;')
    conn.commit()

    # go check all files and wite to database!
    for path in database.iterdir():
        doc = docx.Document(Path(path))
        full_text=[]
        for elem in doc.paragraphs:
            full_text.append(elem.text)

        prediction=get_predictions(our_data.text.values.tolist(),full_text)
        
        result = ""
        for subresult in prediction:
            result += str(subresult) + " ;; "

        try:
            cur.execute("INSERT INTO gosts (name, equip) VALUES (%s, %s) ON CONFLICT DO NOTHING;", (str(path.name).replace('.docx', ''), result))
            conn.commit()
            logger.success(f'Added {path.name}')
        except Exception as e:
            conn.rollback()
            logger.error(e)

    return Response(status_code=200)


@app.post("/predict")
async def predictByName(name: str):
    """
    Функция для определения оборудования в соответсвии с названием (определяет группу продукции, далее госты, далее необзодимое оборудование)

    ДЛЯ РАБОТЫ НЕОБХОДИМО РАСПАКОВАТЬ ПАПКУ models в путь fastApi/names/

    https://disk.yandex.ru/d/kzPozoBRsiJLgQ
    """

    # Example of using the function
    if name:
        input_text = name
        # predicted_class = predict_class(input_text)
        predicted_class = predict_text(name, Path.cwd().joinpath("names").joinpath("models").joinpath("text_classification_pipeline.pkl"))


        labels_keys = list(labels.keys())
        labels_values = list(labels.values()) 

        group = None
        for index in range(len(labels_values)):
            if labels_values[index] == predicted_class:
                group = labels_keys[index]
                    
                cur.execute("SELECT * FROM gosts;")
                data = cur.fetchall()

                total_gost_data = []
                for index, subdata in enumerate(data):
                    # logger.debug(subdata[0]) # id
                    # logger.debug(subdata[1]) # name
                    # logger.debug(subdata[2]) # equip

                    total_gost_data.append({
                        "name" : subdata[1],
                        "equip": subdata[2]
                    })

                # Решение как в первой подзадаче
                uppergroup_on_current_group = []

                # try to find in uppergroups
                for uppergroup in gost_compare:
                    key = list(uppergroup.keys())[0]
                    values = list(uppergroup.values())[0]

                    if group in values:
                        uppergroup_on_current_group.append(key)

                # try to find gost on uppergroups 
                uppergroup_gosts = [] 
                for subgroup in uppergroup_on_current_group:
                    subgroup_gosts = standars_info[standars_info["Группа продукции"] == subgroup]["Обозначение и наименование стандарта"].to_list()
                    uppergroup_gosts.extend(subgroup_gosts)
                uppergroup_gosts = list(set(uppergroup_gosts))

                # check with our table && compare search result to validate equipment
                find_gosts = []
                find_equipment = []

                # if find groups
                if uppergroup_gosts:
                    for subdata in total_gost_data:
                        if subdata["name"].replace(' ', '').replace('"', '').replace("'", '').replace('\xa0', '') in [x.replace(' ', '').replace('"', '').replace("'", '').replace('\xa0', '') for x in uppergroup_gosts]:
                            find_gosts.append(subdata["name"].replace('\xa0', ' '))
                            find_equipment.extend(subdata["equip"].split(';;'))

                find_gosts = list(set(find_gosts))
                find_equipment = list(set(find_equipment))

                logger.success(group)

                return JSONResponse(
                    status_code = 200, 
                    content = {
                        "name": name,
                        "group": group,
                        "find_gosts": find_gosts,
                        "find_equipment": find_equipment
                    }
                )
        
    return Response(status_code=422)