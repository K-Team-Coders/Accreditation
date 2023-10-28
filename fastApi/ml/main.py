import os
import nltk
import pandas as pd
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import linear_kernel
import numpy as np 
from loguru import logger  
import shutil
from nltk.corpus import stopwords
from pymystem3 import Mystem
from string import punctuation

import docx
from pathlib import Path

from fastapi import FastAPI,File,Request,UploadFile

app=FastAPI()

our_data=pd.read_csv(Path.cwd().joinpath("data").joinpath("data_classes.csv"))
mystem = Mystem() 
russian_stopwords = stopwords.words("russian")
nltk.download("stopwords")



def preprocess_text(text):
    tokens = mystem.lemmatize(text.lower())
    tokens = [token for token in tokens if token not in russian_stopwords\
              and token != " " \
              and token.strip() not in punctuation]
    
    text = " ".join(tokens)
    
    return text

def get_predictions(input_strings,target_strings,similarity_threshold=0.5):
    tfidf_vectorizer = TfidfVectorizer(ngram_range=(1, 2))

# Преобразуем тексты в матрицу TF-IDF
    tfidf_matrix_input = tfidf_vectorizer.fit_transform(input_strings)
    tfidf_matrix_target = tfidf_vectorizer.transform(target_strings)

    # Вычисляем косинусное сходство между входными и целевыми строками
    cosine_similarities = linear_kernel(tfidf_matrix_target, tfidf_matrix_input)

    # Находим индексы строк с сходством выше порога
    most_similar_indices = np.argwhere(cosine_similarities > similarity_threshold)

# Выводим только строки, у которых сходство больше 70%
    list_values={}
    for i, j in most_similar_indices:
        input_str = input_strings[j]
        target_str = target_strings[i]
        similarity = cosine_similarities[i, j]
        list_values[input_str]=similarity
        logger.debug(f"Целевая строка: {target_str}")
        logger.debug(f"Наиболее подходящая строка: {input_str}")
        logger.debug(f"Сходство (TF-IDF): {similarity}")
        logger.debug("---------")
    return set(list_values)


@app.post("/get_pred")
async def upload_folder_get_pred(request: Request,  file: UploadFile = File(...)):
    # Get the folder name from the request parameter
    path_to_save=Path.cwd().joinpath('data').joinpath(file.filename)
    with open(path_to_save, "wb") as dest_file:
            shutil.copyfileobj(file.file, dest_file)

    doc = docx.Document(Path(path_to_save))
    full_text=[]
    for elem in doc.paragraphs:
        full_text.append(elem.text)

    prediction=get_predictions(our_data.text.values.tolist(),full_text)

    return {"prediction": prediction}



if __name__=="__main__":
    import uvicorn
    uvicorn.run("main:app", host="localhost", port=8000, reload=True)
    # our_data=pd.read_csv(Path.cwd().joinpath("notebooks\data_chels.csv"))
    # doc = docx.Document(Path.cwd().joinpath("notebooks\ГОСТ Р ИСО 16000-6-2007 Воздух замкнутых помещений. Часть 6....docx"))
    # full_text=[]
    # for elem in doc.paragraphs:
    #     full_text.append(elem.text)
    # get_predictions(our_data.text.values.tolist(),full_text)